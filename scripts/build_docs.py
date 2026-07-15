from __future__ import annotations

import re
import subprocess
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Mm, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "BCube_Prompt_Engine_v3" / "BCube_Prompt_Engine_v3.md"
OUT = ROOT / "docs" / "BCube_Prompt_Engine_v3" / "generated"
TABLE_WIDTH_DXA = 9750


def set_font(run, name="Aptos", size=10.5, color="222222", bold=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def configure(doc):
    section = doc.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(18)
    section.bottom_margin = Mm(18)
    section.left_margin = Mm(19)
    section.right_margin = Mm(19)
    section.header_distance = Mm(10)
    section.footer_distance = Mm(10)
    styles = doc.styles
    tokens = {
        "Normal": (10.5, "222222", False, 0, 6, 1.15),
        "Title": (28, "153C75", True, 0, 8, 1.0),
        "Subtitle": (15, "5271A8", False, 0, 14, 1.0),
        "Heading 1": (16, "153C75", True, 18, 9, 1.05),
        "Heading 2": (13, "2D5A9A", True, 14, 7, 1.05),
        "Heading 3": (11.5, "355C7D", True, 10, 5, 1.05),
    }
    for name, (size, color, bold, before, after, line) in tokens.items():
        st = styles[name]
        st.font.name = "Aptos"
        st._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Aptos")
        st._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Aptos")
        st.font.size = Pt(size)
        st.font.color.rgb = RGBColor.from_string(color)
        st.font.bold = bold
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.line_spacing = line
        st.paragraph_format.keep_with_next = name.startswith("Heading")
    return section


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd")) or OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    if shd.getparent() is None:
        tc_pr.append(shd)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in("w:tcMar")
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tcMar.find(qn(f"w:{m}")) or OxmlElement(f"w:{m}")
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")
        if node.getparent() is None:
            tcMar.append(node)


def set_table_geometry(table, widths):
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_ind.set(qn("w:w"), "120")
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                cell._tc.get_or_add_tcPr().append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(widths[idx]))


def mark_header_row(row):
    tr_pr = row._tr.get_or_add_trPr()
    marker = tr_pr.find(qn("w:tblHeader"))
    if marker is None:
        marker = OxmlElement("w:tblHeader")
        tr_pr.append(marker)
    marker.set(qn("w:val"), "true")


def create_decimal_numbering(doc):
    root = doc.part.numbering_part.element
    abstract_ids = [int(x.get(qn("w:abstractNumId"))) for x in root.findall(qn("w:abstractNum"))]
    num_ids = [int(x.get(qn("w:numId"))) for x in root.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1
    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start"); start.set(qn("w:val"), "1"); level.append(start)
    fmt = OxmlElement("w:numFmt"); fmt.set(qn("w:val"), "decimal"); level.append(fmt)
    text = OxmlElement("w:lvlText"); text.set(qn("w:val"), "%1."); level.append(text)
    suffix = OxmlElement("w:suff"); suffix.set(qn("w:val"), "tab"); level.append(suffix)
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab"); tab.set(qn("w:val"), "num"); tab.set(qn("w:pos"), "540"); tabs.append(tab)
    p_pr.append(tabs)
    ind = OxmlElement("w:ind"); ind.set(qn("w:left"), "540"); ind.set(qn("w:hanging"), "270"); p_pr.append(ind)
    level.append(p_pr)
    abstract.append(level)
    root.append(abstract)
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId"); abstract_ref.set(qn("w:val"), str(abstract_id)); num.append(abstract_ref)
    root.append(num)
    return num_id


def apply_numbering(paragraph, num_id):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = OxmlElement("w:ilvl"); ilvl.set(qn("w:val"), "0"); num_pr.append(ilvl)
    num = OxmlElement("w:numId"); num.set(qn("w:val"), str(num_id)); num_pr.append(num)


def add_page_field(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    set_font(run, size=8.5, color="6B7280")
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    paragraph._p.append(fld)


def add_inline(p, text):
    parts = re.split(r"(`[^`]+`|\*\*[^*]+\*\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("`"):
            r = p.add_run(part[1:-1]); set_font(r, name="Consolas", size=9, color="334155")
        elif part.startswith("**"):
            r = p.add_run(part[2:-2]); set_font(r, bold=True)
        else:
            r = p.add_run(part); set_font(r)


def build():
    text = SOURCE.read_text(encoding="utf-8")
    text = re.sub(r"\A---\n.*?\n---\n", "", text, flags=re.S)
    lines = text.splitlines()
    doc = Document()
    section = configure(doc)
    hp = section.header.paragraphs[0]
    hp.text = "BCUBE FUTURE ACADEMY  |  PROMPT ENGINE v3.0"
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_font(hp.runs[0], size=8.5, color="5271A8", bold=True)
    add_page_field(section.footer.paragraphs[0])

    in_fence = False
    table_rows = []
    first_h1 = True
    current_decimal_num_id = None
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        if line.startswith("```"):
            in_fence = not in_fence
            i += 1
            continue
        if in_fence:
            i += 1
            continue
        if line.startswith("|"):
            table_rows.append([c.strip() for c in line.strip("|").split("|")])
            i += 1
            continue
        elif table_rows:
            rows = [r for r in table_rows if not all(re.fullmatch(r"-+", c) for c in r)]
            if rows:
                cols = max(len(r) for r in rows)
                table = doc.add_table(rows=len(rows), cols=cols)
                table.alignment = WD_TABLE_ALIGNMENT.LEFT
                weights = []
                for ci in range(cols):
                    longest = max(len(row[ci]) if ci < len(row) else 0 for row in rows)
                    weights.append(max(8, min(36, longest)))
                raw_widths = [round(TABLE_WIDTH_DXA * weight / sum(weights)) for weight in weights]
                raw_widths[-1] += TABLE_WIDTH_DXA - sum(raw_widths)
                set_table_geometry(table, raw_widths)
                mark_header_row(table.rows[0])
                for ri, row in enumerate(rows):
                    for ci in range(cols):
                        cell = table.cell(ri, ci)
                        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                        set_cell_margins(cell)
                        if ri == 0:
                            shade(cell, "E8EEF5")
                        p = cell.paragraphs[0]
                        add_inline(p, row[ci] if ci < len(row) else "")
                        for r in p.runs:
                            r.font.size = Pt(8.5)
                            if ri == 0: r.bold = True
                doc.add_paragraph().paragraph_format.space_after = Pt(2)
            table_rows = []
            continue

        if not line.strip():
            i += 1
            continue
        is_decimal_item = bool(re.match(r"\d+\. ", line))
        if not is_decimal_item:
            current_decimal_num_id = None
        if line.startswith("# "):
            if first_h1:
                p = doc.add_paragraph(style="Title")
                add_inline(p, line[2:])
                first_h1 = False
            else:
                doc.add_page_break()
                p = doc.add_paragraph(line[2:], style="Heading 1")
        elif line.startswith("## "):
            style = "Subtitle" if first_h1 is False and len(doc.paragraphs) < 4 else "Heading 2"
            p = doc.add_paragraph(line[3:], style=style)
        elif line.startswith("### "):
            p = doc.add_paragraph(line[4:], style="Heading 3")
        elif line.startswith("#### "):
            p = doc.add_paragraph(line[5:], style="Heading 3")
        elif line.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.left_indent = Inches(0.375)
            p.paragraph_format.first_line_indent = Inches(-0.188)
            p.paragraph_format.space_after = Pt(3)
            spacer = p.add_run("\t")
            set_font(spacer)
            add_inline(p, line[2:])
        elif is_decimal_item:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(3)
            if current_decimal_num_id is None:
                current_decimal_num_id = create_decimal_numbering(doc)
            apply_numbering(p, current_decimal_num_id)
            add_inline(p, re.sub(r"^\d+\. ", "", line))
        elif line.startswith("> "):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.2)
            p.paragraph_format.right_indent = Inches(0.2)
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(8)
            add_inline(p, line[2:])
            for r in p.runs:
                r.italic = True; r.font.color.rgb = RGBColor.from_string("355C7D")
        elif line == "---":
            pass
        else:
            p = doc.add_paragraph()
            add_inline(p, line)
        i += 1

    props = doc.core_properties
    props.title = "BCube Prompt Engine v3.0 - Complete Production Standard"
    props.subject = "Founder Edition"
    props.author = "BCube Future Academy"
    props.comments = "Generated from the canonical Markdown source."
    OUT.mkdir(parents=True, exist_ok=True)
    target = OUT / "BCube_Prompt_Engine_v3.docx"
    doc.save(target)
    print(target)
    return target


if __name__ == "__main__":
    build()
